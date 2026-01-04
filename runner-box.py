# %%
from __future__ import annotations

import json
from dataclasses import dataclass
from typing import Any, Dict, Iterable, List, Mapping, Optional, Sequence, Tuple, Union

from rich.console import Console
from rich.json import JSON
from rich.panel import Panel
from rich.rule import Rule
from rich.syntax import Syntax
from rich.text import Text


# -----------------------------
# Configuration
# -----------------------------

@dataclass(frozen=True)
class Theme:
    user_title: str = "User"
    user_style: str = "bold blue"

    assistant_title: str = "Assistant"
    assistant_style: str = "bold green"

    tool_call_title: str = "Tool Call"
    tool_call_style: str = "bold magenta"

    tool_output_title: str = "Tool Output"
    tool_output_style: str = "bold yellow"

    file_title: str = "File Artifact"
    file_style: str = "bold cyan"

    system_title: str = "System"
    system_style: str = "dim"

    divider_style: str = "dim white"


# -----------------------------
# Renderer
# -----------------------------

class RichAgentRenderer:
    """
    Render streamed events and final outputs from LangChain/LangGraph-style agents
    using Rich panels.

    Design goals:
    - Minimal dependencies on LangGraph internals (e.g., Overwrite wrapper)
    - Robust to multiple message/content formats
    - Deterministic, readable output with safe truncation
    """

    def __init__(
        self,
        console: Optional[Console] = None,
        theme: Theme = Theme(),
        markdown_lexer: str = "markdown",
        syntax_theme: str = "monokai",
        file_preview_lines: int = 999,
    ) -> None:
        self.console = console or Console()
        self.theme = theme
        self.markdown_lexer = markdown_lexer
        self.syntax_theme = syntax_theme
        self.file_preview_lines = file_preview_lines

    # -------------------------
    # Public API
    # -------------------------

    def show_prompt(self, prompt_text: str, title: str = "Prompt", border_style: str = "blue") -> None:
        """
        Render a system or orchestrator prompt in a styled panel.
        """
        text = Text(prompt_text)
        text.highlight_regex(r"^#+.*$", style="bold magenta")
        text.highlight_regex(r"<[^>]+>", style="bold cyan")

        self.console.print(
            Panel(
                text,
                title=f"[bold green]{title}[/bold green]",
                border_style=border_style,
                padding=(1, 2),
            )
        )

    def render_stream_event(self, event: Mapping[str, Any]) -> None:
        """
        Render a single streamed event emitted by an agent.

        Expected shape:
            {"event_type": payload}
        """
        event_type, payload = self._extract_single_kv(event)

        self._divider(event_type)

        if event_type in {"PatchToolCallsMiddleware.before_agent", "model", "tools"}:
            for msg in self._extract_messages(payload):
                self.render_message(msg)

            if event_type == "tools":
                self._render_files_from_payload(payload)

        else:
            # Unknown/middleware events: show payload in JSON form
            self._render_system_payload(payload)

    def render_final_output(self, result: Mapping[str, Any]) -> None:
        """
        Render the final output returned by agent.invoke(...).

        Expected shape:
            {"messages": [...], "files": {...}, ...}
        """
        self._divider("FINAL OUTPUT")

        for msg in self._extract_messages(result):
            self.render_message(msg)

        files = self._get_payload_value(result, "files", default={}) or {}
        if isinstance(files, Mapping) and files:
            self._divider("FILES")
            for path, meta in files.items():
                self._render_file_meta(str(path), meta)

    def render_message(self, message: Any) -> None:
        """
        Render a single LangChain-style message or compatible object.
        """
        cls_name = message.__class__.__name__

        # Support common LangChain message class names
        if cls_name == "HumanMessage":
            self._render_text_panel(
                text=self._format_message_content(message),
                title=self.theme.user_title,
                border_style=self.theme.user_style,
            )
            return

        if cls_name == "AIMessage":
            self._render_text_panel(
                text=self._format_message_content(message) or "",
                title=self.theme.assistant_title,
                border_style=self.theme.assistant_style,
            )
            self._render_tool_calls_from_message(message)
            return

        if cls_name == "ToolMessage":
            self._render_tool_output_message(message)
            return

        # Fallback: render whatever we can
        self._render_text_panel(
            text=self._format_message_content(message),
            title=cls_name,
            border_style="white",
        )

    # -------------------------
    # Message formatting
    # -------------------------

    def _format_message_content(self, message: Any) -> str:
        """
        Extract and render message.content robustly. Supports:
        - Plain strings
        - List-of-blocks with {"type": "text"} / {"type": "tool_use"}
        - Any other object via str(...)
        """
        content = getattr(message, "content", "")

        parts: List[str] = []
        tool_calls_from_content = False

        if isinstance(content, str):
            parts.append(content)

        elif isinstance(content, list):
            for item in content:
                if not isinstance(item, dict):
                    parts.append(str(item))
                    continue

                item_type = item.get("type")
                if item_type == "text":
                    parts.append(str(item.get("text", "")))
                elif item_type in {"tool_use", "tool_call"}:
                    tool_calls_from_content = True
                    name = item.get("name", "unknown_tool")
                    tool_input = item.get("input", item.get("args", {}))
                    call_id = item.get("id", "N/A")
                    parts.append(f"[{self.theme.tool_call_title}] {name}")
                    parts.append(json.dumps(tool_input, indent=2, ensure_ascii=False))
                    parts.append(f"id: {call_id}")
                else:
                    parts.append(str(item))

        else:
            parts.append(str(content))

        # If tool calls were not embedded in content blocks, also check tool_calls attribute
        if not tool_calls_from_content:
            self._append_tool_calls_attribute(parts, message)

        return "\n".join(p for p in parts if p is not None and p != "")

    def _append_tool_calls_attribute(self, parts: List[str], message: Any) -> None:
        tool_calls = getattr(message, "tool_calls", None)
        if not tool_calls:
            return

        # tool_calls usually: [{"name": ..., "args": ..., "id": ...}, ...]
        if isinstance(tool_calls, Sequence):
            for call in tool_calls:
                if not isinstance(call, Mapping):
                    continue
                parts.append(f"[{self.theme.tool_call_title}] {call.get('name', 'unknown_tool')}")
                parts.append(json.dumps(call.get("args", {}), indent=2, ensure_ascii=False))
                parts.append(f"id: {call.get('id', 'N/A')}")

    def _render_tool_calls_from_message(self, message: Any) -> None:
        tool_calls = getattr(message, "tool_calls", None)
        if not tool_calls or not isinstance(tool_calls, Sequence):
            return

        for call in tool_calls:
            if not isinstance(call, Mapping):
                continue
            name = str(call.get("name", "unknown_tool"))
            args = call.get("args", {})
            call_id = call.get("id", "N/A")

            self._render_json_panel(
                data={"name": name, "args": args, "id": call_id},
                title=f"{self.theme.tool_call_title} — {name}",
                border_style=self.theme.tool_call_style,
            )

    def _render_tool_output_message(self, message: Any) -> None:
        content = getattr(message, "content", "")

        # Tool outputs are often JSON strings; parse if possible.
        if isinstance(content, str):
            parsed = self._try_parse_json(content)
            if parsed is not None:
                self._render_json_panel(
                    data=parsed,
                    title=self.theme.tool_output_title,
                    border_style=self.theme.tool_output_style,
                )
                return

        self._render_text_panel(
            text=str(content),
            title=self.theme.tool_output_title,
            border_style=self.theme.tool_output_style,
        )

    # -------------------------
    # Stream payload helpers
    # -------------------------

    def _extract_messages(self, payload: Any) -> List[Any]:
        """
        Extract 'messages' from dict-like payloads; returns an empty list if missing.
        Normalizes singleton -> list.
        """
        payload = self._unwrap_overwrite(payload)

        messages = self._get_payload_value(payload, "messages", default=[])
        messages = self._unwrap_overwrite(messages)

        if messages is None:
            return []
        if isinstance(messages, (list, tuple)):
            return list(messages)
        return [messages]

    def _render_files_from_payload(self, payload: Any) -> None:
        payload = self._unwrap_overwrite(payload)
        files = self._get_payload_value(payload, "files", default={}) or {}
        if not isinstance(files, Mapping):
            return
        for path, meta in files.items():
            self._render_file_meta(str(path), meta)

    def _render_file_meta(self, path: str, meta: Any) -> None:
        meta = meta or {}
        if not isinstance(meta, Mapping):
            meta = {"meta": str(meta)}

        preview = meta.get("content", [])
        if isinstance(preview, list):
            preview = preview[: self.file_preview_lines]
        else:
            preview = str(preview)

        self._render_json_panel(
            data={
                "path": path,
                "created_at": meta.get("created_at"),
                "modified_at": meta.get("modified_at"),
                "preview": preview,
            },
            title=f"{self.theme.file_title} — {path}",
            border_style=self.theme.file_style,
        )

    def _render_system_payload(self, payload: Any) -> None:
        payload = self._unwrap_overwrite(payload)

        # Prefer JSON panel for dict/list payloads
        if isinstance(payload, (dict, list)):
            self._render_json_panel(payload, self.theme.system_title, self.theme.system_style)
        else:
            self._render_text_panel(str(payload), self.theme.system_title, self.theme.system_style)

    # -------------------------
    # Low-level rendering
    # -------------------------

    def _divider(self, label: str = "") -> None:
        self.console.print(Rule(label, style=self.theme.divider_style))

    def _render_text_panel(self, text: str, title: str, border_style: str) -> None:
        syntax = Syntax(
            text or "",
            lexer=self.markdown_lexer,
            theme=self.syntax_theme,
            word_wrap=True,
        )
        self.console.print(
            Panel(
                syntax,
                title=title,
                border_style=border_style,
                padding=(1, 2),
            )
        )

    def _render_json_panel(self, data: Any, title: str, border_style: str) -> None:
        self.console.print(
            Panel(
                JSON.from_data(data, indent=2),
                title=title,
                border_style=border_style,
                padding=(1, 1),
            )
        )

    # -------------------------
    # Utilities
    # -------------------------

    @staticmethod
    def _extract_single_kv(mapping: Mapping[str, Any]) -> Tuple[str, Any]:
        if not mapping:
            raise ValueError("Event payload is empty.")
        if len(mapping) != 1:
            # Keep deterministic behavior, but avoid silent misrender
            raise ValueError(f"Expected exactly 1 top-level key in event, got {len(mapping)}.")
        return next(iter(mapping.items()))

    @staticmethod
    def _try_parse_json(text: str) -> Optional[Any]:
        try:
            return json.loads(text)
        except Exception:
            return None

    @staticmethod
    def _unwrap_overwrite(value: Any) -> Any:
        """
        Unwrap LangGraph Overwrite wrapper without requiring a direct import.
        """
        if value is None:
            return None
        # Common pattern: Overwrite(value=<payload>)
        if getattr(value, "__class__", None) and value.__class__.__name__ == "Overwrite":
            return getattr(value, "value", value)
        return value

    @staticmethod
    def _get_payload_value(payload: Any, key: str, default: Any = None) -> Any:
        payload = RichAgentRenderer._unwrap_overwrite(payload)
        if isinstance(payload, Mapping):
            return payload.get(key, default)
        return getattr(payload, key, default)


# -----------------------------
# Convenience functions
# -----------------------------

_DEFAULT_RENDERER = RichAgentRenderer()

_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialized RichAgentRenderer for rendering LangChain agent events and outputs."),
        title="STEP 1: Initialization",
        border_style="purple",
        padding=(1, 2),
    )
)

#%%

_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize your environment by loading .env variables."),
        title="STEP 2: Initialize Environment",
        border_style="purple",
        padding=(1, 2),
    )
)
import os
from typing import Any
from dotenv import load_dotenv

load_dotenv(".env", override=True)

# %%
_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize tools and affiliated resources."),
        title="STEP 3: Initialize Tools",
        border_style="purple",
        padding=(1, 2),
    )
)
from pathlib import Path
from langchain.tools import tool
from langgraph.types import Overwrite

from pathlib import Path

WORKSPACE_ROOT = Path("./workspace").resolve()


def resolve_workspace_path(virtual_path: str) -> Path:
    """
    Resolve a virtual workspace path into a real filesystem path.

    Accepts:
    - /workspace
    - /workspace/
    - /workspace/relative/path

    Rejects:
    - Any path outside the workspace namespace
    - Any path attempting directory traversal
    """

    if virtual_path == "/workspace":
        return WORKSPACE_ROOT

    if virtual_path.startswith("/workspace/"):
        relative = virtual_path[len("/workspace/") :]
        real_path = (WORKSPACE_ROOT / relative).resolve()
    else:
        raise ValueError(f"Invalid workspace path: {virtual_path}")

    # Enforce sandboxing
    if not str(real_path).startswith(str(WORKSPACE_ROOT)):
        raise ValueError(f"Path traversal detected: {virtual_path}")

    return real_path


def safe_fix_zip_filename(name: str) -> str:
    """
    Attempt to fix garbled ZIP filenames produced by legacy tools.
    Never raises UnicodeEncodeError.
    """
    try:
        raw = name.encode("latin1")  # latin1 is always reversible
    except Exception:
        return name

    for enc in ("utf-8", "gbk", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue

    return name


#%%
_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize unzip tool."),
        title="STEP 3: Initialize Tools",
        border_style="purple",
        padding=(1, 2),
    )
)
import zipfile


@tool(parse_docstring=True)
def unzip_workspace_file(virtual_zip_path: str) -> dict:
    """
    Unzip a ZIP archive located inside the workspace.

    This preprocessing function extracts the contents of a ZIP file referenced
    by a virtual workspace path. The archive is unpacked into a directory with
    the same base name as the ZIP file, located in the same workspace directory.

    The function operates strictly within the workspace sandbox:
    - Input paths are virtual (e.g. `/workspace/foo.zip`)
    - Execution is performed on resolved real paths
    - Output paths are returned in virtual-path form for agent consumption

    Typical agent usage:
    - Call this when a ZIP file is detected during workspace inspection
    - Follow with a tree-view operation on the extracted directory

    Args:
        virtual_zip_path: Virtual path to a ZIP file inside the workspace
            (for example `/workspace/datasets/images.zip`). PATH MUST START WITH `/workspace`.

    Returns:
        A dictionary containing:
        - status: Execution status string
        - zip: The input virtual ZIP path
        - extracted_to: Virtual path of the extracted directory
        - num_files: Number of files listed in the ZIP archive

    Raises:
        FileNotFoundError: If the ZIP file does not exist.
        ValueError: If the provided path does not point to a ZIP archive.
    """

    # zip_path = resolve_workspace_path(virtual_zip_path)

    # if not zip_path.exists():
    #     raise FileNotFoundError(zip_path)

    # if zip_path.suffix.lower() != ".zip":
    #     raise ValueError("Provided file is not a zip archive")

    # output_dir = zip_path.with_suffix("")
    # output_dir.mkdir(parents=True, exist_ok=True)

    # with zipfile.ZipFile(zip_path, "r") as zf:
    #     zf.extractall(output_dir)

    # return {
    #     "status": "ok",
    #     "zip": virtual_zip_path,
    #     "extracted_to": f"/workspace/{output_dir.name}",
    #     "num_files": len(zf.namelist()),
    # }

    zip_path = resolve_workspace_path(virtual_zip_path)

    if not zip_path.exists():
        raise FileNotFoundError(zip_path)

    if zip_path.suffix.lower() != ".zip":
        raise ValueError("Provided file is not a zip archive")

    output_dir = zip_path.with_suffix("")
    output_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(zip_path, "r") as zf:
        infos = zf.infolist()
        for info in infos:
            info.filename = safe_fix_zip_filename(info.filename)
            zf.extract(info, output_dir)

    return {
        "status": "ok",
        "zip": virtual_zip_path,
        "extracted_to": f"/workspace/{output_dir.name}",
        "num_files": len(infos),
    }


# %%
_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize tree view tool."),
        title="STEP 3: Initialize Tools",
        border_style="purple",
        padding=(1, 2),
    )
)
@tool(parse_docstring=True)
def tree_view_workspace(
    virtual_path: str,
    max_depth: int = 4,
    max_entries: int = 200,
) -> dict:
    """
    Generate a hierarchical tree view of a workspace directory.

    This preprocessing function inspects the directory structure rooted at a
    given virtual workspace path and produces a readable tree representation.
    It is designed for lightweight structural understanding rather than full
    content analysis.

    The function supports depth and entry limits to prevent excessive output
    and is suitable for agent-driven exploration and reporting.

    Typical agent usage:
    - Inspect workspace contents before deciding on further preprocessing
    - Summarize directory layout in reports such as `/final_report.md`
    - Validate results of prior operations such as unzip or file generation

    Args:
        virtual_path: Virtual workspace path to inspect, such as `/workspace`
            or `/workspace/data` PATH MUST START WITH `/workspace`.
        max_depth: Maximum directory depth to traverse. The default value limits
            traversal to a small number of levels to avoid excessive output.
        max_entries: Maximum total number of files or directories to include
            in the output. The default value prevents large directories from
            producing oversized results.

    Returns:
        A dictionary containing:
        - root: The inspected virtual path
        - max_depth: The depth limit used during traversal
        - entries: Number of directory entries included
        - tree: A newline-separated string representing the directory tree
        - truncated: Whether traversal stopped early due to entry limits

    Raises:
        FileNotFoundError: If the target path does not exist.
    """

    root = resolve_workspace_path(virtual_path)

    if not root.exists():
        raise FileNotFoundError(root)

    lines = []
    count = 0

    def walk(path: Path, prefix: str = "", depth: int = 0):
        nonlocal count
        if depth > max_depth or count >= max_entries:
            return

        for p in sorted(path.iterdir()):
            if count >= max_entries:
                return

            lines.append(f"{prefix}{p.name}")
            count += 1

            if p.is_dir():
                walk(p, prefix + "  ", depth + 1)

    walk(root)

    return {
        "root": virtual_path,
        "max_depth": max_depth,
        "entries": count,
        "tree": "\n".join(lines),
        "truncated": count >= max_entries,
    }


# %%

_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize PDF reader tool."),
        title="STEP 3: Initialize Tools",
        border_style="purple",
        padding=(1, 2),
    )
)
@tool(parse_docstring=True)
def pdf_reader(virtual_pdf_path: str, num_pages: int = 5) -> dict:
    """
    Read the first N pages of a PDF file located inside the workspace.

    This preprocessing function extracts text from the first N pages of a PDF
    file referenced by a virtual workspace path.

    The function operates strictly within the workspace sandbox:
    - Input paths are virtual (e.g. `/workspace/document.pdf`)
    - Execution is performed on resolved real paths
    - Output paths are returned in virtual-path form for agent consumption

    Typical agent usage:
    - Call this when a PDF file is detected during workspace inspection
    - Follow with further operations based on the extracted text

    Args:
        virtual_pdf_path: Virtual path to a PDF file inside the workspace
            (for example `/workspace/document.pdf`). PATH MUST START WITH `/workspace`.
        num_pages: Number of pages to read from the start of the PDF.
    
    Returns:
        A dictionary containing:
        - status: Execution status string
        - pdf: The input virtual PDF path
        - content: Extracted text from the first N pages of the PDF
    
    Raises:
        FileNotFoundError: If the PDF file does not exist.
        ValueError: If the provided path does not point to a PDF file.
    """
    from pypdf import PdfReader

    pdf_path = resolve_workspace_path(virtual_pdf_path)

    if not pdf_path.exists():
        return {
            "status": "error",
            "content": f"File not found: {virtual_pdf_path}"
        }

    if pdf_path.suffix.lower() != ".pdf":
        raise ValueError("Provided file is not a PDF file")

    reader = PdfReader(str(pdf_path))
    content = []

    for i, page in enumerate(reader.pages):
        if i >= num_pages:
            break
        content.append(page.extract_text())

    return {
        "status": "ok",
        "pdf": virtual_pdf_path,
        "content": "\n".join(content),
    }
    
# %%
_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize Word reader tool."),
        title="STEP 3: Initialize Tools",
        border_style="purple",
        padding=(1, 2),
    )
)
@tool(parse_docstring=True)
def word_reader(virtual_docx_path: str, max_blocks: int = 200) -> dict:
    """
    Read a Microsoft Word document and convert it to formatted Markdown.

    This preprocessing function extracts structured content from a Word document
    referenced by a virtual workspace path. It returns a Markdown rendering plus
    formatting metadata for paragraphs, runs, and tables.

    The function operates strictly within the workspace sandbox:
    - Input paths are virtual (e.g. `/workspace/document.docx`)
    - Execution is performed on resolved real paths
    - No file content is modified

    Typical agent usage:
    - Call this when a Word document is detected during workspace inspection
    - Use the Markdown output for downstream summarization or extraction
    - Use the formatting metadata when layout fidelity matters

    Args:
        virtual_docx_path: Virtual path to a Word document inside the workspace
            (for example `/workspace/docs/report.docx`).
            PATH MUST START WITH `/workspace`.
        max_blocks: Maximum number of top-level blocks (paragraphs or tables)
            to return. Set to 0 to disable the limit.

    Returns:
        A dictionary containing:
        - status: Execution status string
        - word: The input virtual Word path
        - metadata: Core document metadata (title, author, timestamps)
        - blocks: Structured blocks with formatting info
        - markdown: Markdown rendering of the document
        - blocks_returned: Number of blocks returned
        - truncated: Whether the output was truncated by max_blocks

    Raises:
        ValueError: If the provided path is not a supported Word document.
    """
    try:
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except Exception as exc:
        return {
            "status": "error",
            "word": virtual_docx_path,
            "error": f"python-docx is required to read Word files: {exc}",
        }

    import re

    docx_path = resolve_workspace_path(virtual_docx_path)

    if not docx_path.exists():
        return {
            "status": "error",
            "word": virtual_docx_path,
            "error": f"File not found: {virtual_docx_path}",
        }

    if docx_path.suffix.lower() not in {".docx", ".docm", ".dotx", ".dotm"}:
        raise ValueError("Provided file is not a supported Word document")

    document = Document(str(docx_path))

    def iter_block_items(doc: Document) -> Iterable[Union[Paragraph, Table]]:
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    def escape_markdown_chars(text: str) -> str:
        text = text.replace("\\", "\\\\")
        for ch in ("`", "*", "_", "{", "}", "[", "]", "(", ")", "#", "+", "!", "|", ">"):
            text = text.replace(ch, f"\\{ch}")
        return text

    def escape_html(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def normalize_text(text: str, in_table: bool = False) -> str:
        if not text:
            return ""
        text = text.replace("\r", "")
        text = text.replace("\t", "    ")
        token = "__DOCX_BR__"
        text = text.replace("\n", token)
        text = escape_html(text)
        text = escape_markdown_chars(text)
        if in_table:
            return text.replace(token, "<br>")
        return text.replace(token, "  \n")

    def resolve_emphasis(run) -> Tuple[bool, bool]:
        bold = run.bold
        italic = run.italic
        style_name = (run.style.name if run.style else "").lower()
        if bold is None and ("strong" in style_name or "bold" in style_name):
            bold = True
        if italic is None and ("emphasis" in style_name or "italic" in style_name):
            italic = True
        return bool(bold), bool(italic)

    def run_to_markdown(run, in_table: bool = False) -> str:
        raw_text = run.text
        if raw_text is None or raw_text == "":
            return ""
        text = normalize_text(raw_text, in_table=in_table)
        bold, italic = resolve_emphasis(run)
        strike = True if run.font.strike else False
        if bold and italic:
            text = f"***{text}***"
        elif bold:
            text = f"**{text}**"
        elif italic:
            text = f"*{text}*"
        if strike:
            text = f"~~{text}~~"
        return text

    def paragraph_to_markdown(paragraph: Paragraph, in_table: bool = False) -> str:
        parts: List[str] = []
        for run in paragraph.runs:
            part = run_to_markdown(run, in_table=in_table)
            if part:
                parts.append(part)
        if not parts:
            return normalize_text(paragraph.text or "", in_table=in_table)
        return "".join(parts)

    def detect_heading_level(style_id: Optional[str], style_name: Optional[str]) -> Optional[int]:
        for source in (style_id, style_name):
            if not source:
                continue
            lower = source.lower()
            if lower.startswith("heading"):
                match = re.search(r"(\d+)", source)
                if match:
                    level = int(match.group(1))
                    return max(1, min(level, 6))
        if style_id and style_id.lower() == "title":
            return 1
        if style_name and style_name.lower() == "title":
            return 1
        if style_id and style_id.lower() == "subtitle":
            return 2
        if style_name and style_name.lower() == "subtitle":
            return 2
        return None

    def detect_list_info(paragraph: Paragraph, style_id: Optional[str], style_name: Optional[str]) -> Tuple[Optional[str], Optional[int]]:
        list_type = None
        list_level = None
        combined = f"{style_id or ''} {style_name or ''}".lower()

        if "bullet" in combined:
            list_type = "bullet"
        elif "number" in combined or "decimal" in combined:
            list_type = "number"
        elif "list" in combined:
            list_type = "number"

        p_pr = paragraph._p.pPr
        if p_pr is not None and p_pr.numPr is not None:
            ilvl = p_pr.numPr.ilvl
            if ilvl is not None and ilvl.val is not None:
                list_level = int(ilvl.val) + 1
            if list_type is None:
                list_type = "number"

        if list_level is None and style_name:
            match = re.search(r"(\d+)$", style_name)
            if match:
                list_level = int(match.group(1))

        if list_type and list_level is None:
            list_level = 1

        return list_type, list_level

    def is_quote_style(style_id: Optional[str], style_name: Optional[str]) -> bool:
        combined = f"{style_id or ''} {style_name or ''}".lower()
        return "quote" in combined

    def extract_run_format(run) -> Dict[str, Any]:
        underline = run.underline
        if underline is not None and underline not in (True, False):
            underline = str(underline)

        color = None
        if run.font.color is not None:
            if run.font.color.rgb is not None:
                color = f"#{run.font.color.rgb}"
            elif run.font.color.theme_color is not None:
                color = str(run.font.color.theme_color)

        highlight = None
        if run.font.highlight_color is not None:
            highlight = str(run.font.highlight_color)

        size_pt = run.font.size.pt if run.font.size else None

        return {
            "bold": run.bold,
            "italic": run.italic,
            "underline": underline,
            "strike": run.font.strike,
            "superscript": run.font.superscript,
            "subscript": run.font.subscript,
            "font_name": run.font.name,
            "font_size_pt": size_pt,
            "color": color,
            "highlight": highlight,
            "style": run.style.name if run.style else None,
        }

    def extract_paragraph_info(paragraph: Paragraph) -> Dict[str, Any]:
        style = paragraph.style
        style_name = style.name if style else None
        style_id = style.style_id if style else None
        heading_level = detect_heading_level(style_id, style_name)
        list_type, list_level = detect_list_info(paragraph, style_id, style_name)
        quote = is_quote_style(style_id, style_name)

        block_type = "paragraph"
        if heading_level:
            block_type = "heading"
        elif list_type:
            block_type = "list_item"
        elif quote:
            block_type = "quote"

        runs: List[Dict[str, Any]] = []
        for run in paragraph.runs:
            if run.text is None or run.text == "":
                continue
            runs.append({"text": run.text, "formatting": extract_run_format(run)})

        alignment = None
        if paragraph.alignment is not None:
            alignment = str(paragraph.alignment)

        return {
            "type": block_type,
            "text": paragraph.text or "",
            "style": style_name,
            "style_id": style_id,
            "alignment": alignment,
            "list_type": list_type,
            "list_level": list_level,
            "heading_level": heading_level,
            "runs": runs,
        }

    def table_to_markdown(table: Table) -> Tuple[str, List[List[str]]]:
        rows: List[List[str]] = []
        for row in table.rows:
            row_cells: List[str] = []
            for cell in row.cells:
                paragraphs = []
                for paragraph in cell.paragraphs:
                    text = paragraph_to_markdown(paragraph, in_table=True).strip()
                    if text:
                        paragraphs.append(text)
                cell_text = "<br>".join(paragraphs) if paragraphs else ""
                row_cells.append(cell_text)
            rows.append(row_cells)

        if not rows:
            return "", []

        col_count = max(len(row) for row in rows)
        for row in rows:
            if len(row) < col_count:
                row.extend([""] * (col_count - len(row)))

        header = rows[0]
        separator = ["---"] * col_count
        body_rows = rows[1:]

        def format_row(values: List[str]) -> str:
            return "| " + " | ".join(values) + " |"

        markdown_lines = [
            format_row(header),
            format_row(separator),
        ]
        for row in body_rows:
            markdown_lines.append(format_row(row))
        return "\n".join(markdown_lines), rows

    def extract_table_info(table: Table) -> Dict[str, Any]:
        rows_info: List[List[Dict[str, Any]]] = []
        for row in table.rows:
            row_info: List[Dict[str, Any]] = []
            for cell in row.cells:
                cell_paragraphs = [extract_paragraph_info(p) for p in cell.paragraphs]
                cell_text = "\n".join(p["text"] for p in cell_paragraphs if p["text"])
                row_info.append(
                    {
                        "text": cell_text,
                        "paragraphs": cell_paragraphs,
                    }
                )
            rows_info.append(row_info)

        column_count = 0
        if table.rows:
            column_count = max(len(row.cells) for row in table.rows)

        return {
            "type": "table",
            "row_count": len(table.rows),
            "column_count": column_count,
            "rows": rows_info,
        }

    metadata = {
        "title": document.core_properties.title,
        "subject": document.core_properties.subject,
        "author": document.core_properties.author,
        "created": document.core_properties.created.isoformat() if document.core_properties.created else None,
        "modified": document.core_properties.modified.isoformat() if document.core_properties.modified else None,
        "last_modified_by": document.core_properties.last_modified_by,
        "revision": document.core_properties.revision,
    }

    blocks: List[Dict[str, Any]] = []
    markdown_blocks: List[str] = []
    list_buffer: List[str] = []
    current_list_type: Optional[str] = None
    truncated = False

    def flush_list_buffer() -> None:
        nonlocal current_list_type
        if list_buffer:
            markdown_blocks.append("\n".join(list_buffer))
            list_buffer.clear()
        current_list_type = None

    block_limit = max_blocks if max_blocks and max_blocks > 0 else None

    for block in iter_block_items(document):
        if block_limit is not None and len(blocks) >= block_limit:
            truncated = True
            break

        if isinstance(block, Paragraph):
            info = extract_paragraph_info(block)
            blocks.append(info)

            style = block.style
            style_name = style.name if style else None
            style_id = style.style_id if style else None
            heading_level = detect_heading_level(style_id, style_name)
            list_type, list_level = detect_list_info(block, style_id, style_name)
            is_quote = is_quote_style(style_id, style_name)
            text_md = paragraph_to_markdown(block).strip()

            if list_type:
                if current_list_type and list_type != current_list_type:
                    flush_list_buffer()
                current_list_type = list_type
                indent = "  " * ((list_level or 1) - 1)
                prefix = "- " if list_type == "bullet" else "1. "
                list_buffer.append(f"{indent}{prefix}{text_md}")
                continue

            flush_list_buffer()

            if heading_level:
                markdown_blocks.append(f"{'#' * heading_level} {text_md}".rstrip())
            elif is_quote:
                quote_lines = text_md.splitlines() or [""]
                markdown_blocks.append("\n".join(f"> {line}".rstrip() for line in quote_lines))
            else:
                markdown_blocks.append(text_md)

        elif isinstance(block, Table):
            flush_list_buffer()
            table_markdown, _ = table_to_markdown(block)
            if table_markdown:
                markdown_blocks.append(table_markdown)
            blocks.append(extract_table_info(block))

    flush_list_buffer()

    markdown_output = "\n\n".join(markdown_blocks).strip()

    return {
        "status": "ok",
        "word": virtual_docx_path,
        "metadata": metadata,
        "blocks": blocks,
        "markdown": markdown_output,
        "blocks_returned": len(blocks),
        "truncated": truncated,
    }

# %%
_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize audio transcription tool."),
        title="STEP 3: Initialize Tools",
        border_style="purple",
        padding=(1, 2),
    )
)
@tool(parse_docstring=True)
def audio_transcribe(
    virtual_audio_path: str,
    server_url: str = "http://127.0.0.1:8080/inference",
    response_format: str = "json",
    temperature: float = 0.0,
    temperature_inc: float = 0.2,
    language: Optional[str] = None,
    prompt: Optional[str] = None,
    convert_to_wav: bool = True,
    timeout_sec: float = 300.0,
) -> dict:
    """
    Transcribe an audio file located inside the workspace using whisper.cpp server.

    This preprocessing function sends a local audio file to the whisper.cpp
    HTTP server and returns the transcription results.

    The function operates strictly within the workspace sandbox:
    - Input paths are virtual (e.g. `/workspace/audio.m4a`)
    - Execution is performed on resolved real paths
    - No file content is modified

    Notes:
    - If the input is not a WAV file, this tool can convert it locally using
      ffmpeg (requires ffmpeg on PATH). Disable with convert_to_wav=False to
      rely on server-side --convert or pre-converted WAV input.

    Typical agent usage:
    - Call this when an audio file is detected during workspace inspection
    - Follow with summarization or downstream extraction tasks

    Args:
        virtual_audio_path: Virtual path to an audio file inside the workspace
            (for example `/workspace/recordings/meeting.m4a`).
            PATH MUST START WITH `/workspace`.
        server_url: Whisper server inference endpoint URL.
        response_format: Response format requested from whisper server.
        temperature: Decode temperature.
        temperature_inc: Temperature increment for fallback decoding.
        language: Spoken language code (e.g. en, zh). Use None for auto-detect.
        prompt: Optional initial prompt to bias the transcription.
        convert_to_wav: Convert non-WAV input to 16 kHz mono PCM WAV locally.
        timeout_sec: HTTP request timeout in seconds.

    Returns:
        A dictionary containing:
        - status: Execution status string
        - audio: The input virtual audio path
        - text: Transcribed text if available
        - segments: Segment list if available
        - response_format: The response format used

    Raises:
        FileNotFoundError: If the audio file does not exist.
        ValueError: If the provided path is outside the workspace.
    """

    import httpx
    import shutil
    import subprocess
    import tempfile

    audio_path = resolve_workspace_path(virtual_audio_path)

    if not audio_path.exists():
        return {
            "status": "error",
            "audio": virtual_audio_path,
            "error": f"File not found: {virtual_audio_path}",
        }

    if not audio_path.is_file():
        return {
            "status": "error",
            "audio": virtual_audio_path,
            "error": "Path is not a file.",
        }

    temp_dir: Optional[tempfile.TemporaryDirectory] = None
    upload_path = audio_path
    converted = False

    if audio_path.suffix.lower() != ".wav":
        if not convert_to_wav:
            return {
                "status": "error",
                "audio": virtual_audio_path,
                "error": "Input is not WAV. Enable convert_to_wav or start the server with --convert.",
            }

        ffmpeg = shutil.which("ffmpeg")
        if not ffmpeg:
            return {
                "status": "error",
                "audio": virtual_audio_path,
                "error": "ffmpeg not found on PATH. Install ffmpeg or disable convert_to_wav.",
            }

        temp_dir = tempfile.TemporaryDirectory()
        output_path = Path(temp_dir.name) / f"{audio_path.stem}.wav"
        cmd = [
            ffmpeg,
            "-hide_banner",
            "-loglevel",
            "error",
            "-y",
            "-i",
            str(audio_path),
            "-ar",
            "16000",
            "-ac",
            "1",
            "-c:a",
            "pcm_s16le",
            str(output_path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, check=False)
        if result.returncode != 0:
            temp_dir.cleanup()
            return {
                "status": "error",
                "audio": virtual_audio_path,
                "error": "ffmpeg failed to convert the input to WAV.",
                "details": result.stderr.strip() or "No stderr output.",
            }

        upload_path = output_path
        converted = True

    data = {
        "temperature": str(temperature),
        "temperature_inc": str(temperature_inc),
        "response_format": response_format,
    }
    if language:
        data["language"] = language
    if prompt:
        data["prompt"] = prompt

    content_type = "audio/wav" if upload_path.suffix.lower() == ".wav" else "application/octet-stream"
    try:
        with upload_path.open("rb") as handle:
            files = {"file": (upload_path.name, handle, content_type)}
            with httpx.Client(timeout=timeout_sec) as client:
                response = client.post(server_url, data=data, files=files)
    except httpx.RequestError as exc:
        return {
            "status": "error",
            "audio": virtual_audio_path,
            "error": f"Request failed: {exc}",
        }
    finally:
        if temp_dir is not None:
            temp_dir.cleanup()

    if response.is_error:
        body = response.text.strip()
        return {
            "status": "error",
            "audio": virtual_audio_path,
            "status_code": response.status_code,
            "reason": response.reason_phrase,
            "error": body or "Server returned an error response.",
        }

    payload: Any
    if response.headers.get("content-type", "").startswith("application/json"):
        payload = response.json()
    else:
        try:
            payload = response.json()
        except json.JSONDecodeError:
            payload = response.text

    if isinstance(payload, str):
        return {
            "status": "ok",
            "audio": virtual_audio_path,
            "converted": converted,
            "response_format": response_format,
            "text": payload.strip(),
        }

    result = payload
    if isinstance(payload, dict) and isinstance(payload.get("result"), dict):
        result = payload["result"]

    if isinstance(result, dict):
        return {
            "status": "ok",
            "audio": virtual_audio_path,
            "converted": converted,
            "response_format": response_format,
            "text": (result.get("text") or "").strip(),
            "segments": result.get("segments", []),
        }

    return {
        "status": "ok",
        "audio": virtual_audio_path,
        "converted": converted,
        "response_format": response_format,
        "response": payload,
    }

# %%
from typing import Dict, Any
import pandas as pd
from langchain.tools import tool


@tool(parse_docstring=True)
def excel_schema_reader(virtual_excel_path: str) -> dict:
    """
    Inspect an Excel file located inside the workspace and return its schema.

    This preprocessing function reads an Excel file referenced by a virtual
    workspace path and extracts high-level structural information, including
    sheet names, row counts, and column metadata.

    The function operates strictly within the workspace sandbox:
    - Input paths are virtual (e.g. `/workspace/data.xlsx`)
    - Execution is performed on resolved real paths
    - No file content is modified

    Typical agent usage:
    - Call this before extracting rows from an Excel file
    - Use the schema to decide which sheet and columns to target
    - Provide structural summaries in reports

    Args:
        virtual_excel_path: Virtual path to an Excel file inside the workspace
            (for example `/workspace/tables/data.xlsx`).
            PATH MUST START WITH `/workspace`.

    Returns:
        A dictionary containing:
        - status: Execution status string
        - excel: The input virtual Excel path
        - sheets: Mapping from sheet name to schema information:
            - num_rows: Number of rows in the sheet
            - columns: List of column descriptors with name and dtype

    Raises:
        FileNotFoundError: If the Excel file does not exist.
        ValueError: If the file is not a valid Excel document.
    """

    excel_path = resolve_workspace_path(virtual_excel_path)

    if not excel_path.exists():
        raise FileNotFoundError(excel_path)

    try:
        sheets = pd.read_excel(excel_path, sheet_name=None)
    except Exception as e:
        raise ValueError(f"Failed to read Excel file: {e}")

    schema = {}

    for sheet_name, df in sheets.items():
        schema[sheet_name] = {
            "num_rows": len(df),
            "columns": [
                {
                    "name": col,
                    "dtype": str(df[col].dtype),
                }
                for col in df.columns
            ],
        }

    return {
        "status": "ok",
        "excel": virtual_excel_path,
        "sheets": schema,
    }
# %%

from typing import Optional, List, Dict, Any


@tool(parse_docstring=True)
def excel_entry_extractor(
    virtual_excel_path: str,
    sheet_name: str,
    columns: Optional[List[str]] = None,
    max_rows: int = 50,
    filters: Optional[Dict[str, Any]] = None,
) -> dict:
    """
    Extract structured row entries from a specific sheet of an Excel file.

    This preprocessing function loads a specified sheet from an Excel file
    located inside the workspace and returns row-level data as structured
    dictionaries. Optional column selection and exact-match filtering are
    supported.

    The function operates strictly within the workspace sandbox:
    - Input paths are virtual (e.g. `/workspace/data.xlsx`)
    - Execution is performed on resolved real paths
    - Returned data is JSON-serializable for agent consumption

    Typical agent usage:
    - Extract tabular records for reasoning or summarization
    - Retrieve example rows for inspection
    - Perform simple rule-based filtering before LLM analysis

    Args:
        virtual_excel_path: Virtual path to an Excel file inside the workspace
            (for example `/workspace/tables/data.xlsx`).
            PATH MUST START WITH `/workspace`.
        sheet_name: Name of the Excel sheet to extract rows from.
        columns: Optional list of column names to return.
            If omitted, all columns are included.
        max_rows: Maximum number of rows to return.
        filters: Optional exact-match filters in the form `{column_name: value}`. Filters are applied sequentially.

    Returns:
        A dictionary containing:
        - status: Execution status string
        - excel: The input virtual Excel path
        - sheet: The targeted sheet name
        - rows_returned: Number of rows returned
        - entries: List of row dictionaries

    """

    excel_path = resolve_workspace_path(virtual_excel_path)

    if not excel_path.exists():
        raise FileNotFoundError(excel_path)

    sheets = pd.read_excel(excel_path, sheet_name=None)

    if sheet_name not in sheets:
        raise ValueError(f"Sheet not found: {sheet_name}")

    df = sheets[sheet_name]

    if filters:
        for col, val in filters.items():
            if col in df.columns:
                df = df[df[col] == val]

    if columns:
        df = df[columns]

    df = df.head(max_rows)

    entries = df.to_dict(orient="records")

    return {
        "status": "ok",
        "excel": virtual_excel_path,
        "sheet": sheet_name,
        "rows_returned": len(entries),
        "entries": entries,
    }
#%%

@tool(parse_docstring=True)
def think_tool(reflection: str) -> str:
    """Tool for strategic reflection on research progress and decision-making.

    Use this tool after each search to analyze results and plan next steps systematically.
    This creates a deliberate pause in the research workflow for quality decision-making.

    When to use:
    - After receiving search results: What key information did I find?
    - Before deciding next steps: Do I have enough to answer comprehensively?
    - When assessing research gaps: What specific information am I still missing?
    - Before concluding research: Can I provide a complete answer now?

    Reflection should address:
    1. Analysis of current findings - What concrete information have I gathered?
    2. Gap assessment - What crucial information is still missing?
    3. Quality evaluation - Do I have sufficient evidence/examples for a good answer?
    4. Strategic decision - Should I continue searching or provide my answer?

    Args:
        reflection: Your detailed reflection on research progress, findings, gaps, and next steps

    Returns:
        Confirmation that reflection was recorded for decision-making
    """
    return f"Reflection recorded: {reflection}"
# %%
all_tools = [
    tree_view_workspace,
    unzip_workspace_file,
    pdf_reader,
    word_reader,
    audio_transcribe,
    excel_entry_extractor,
    excel_schema_reader,
    think_tool
]
# %%
_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize Prompts."),
        title="STEP 4: Initialize Prompts",
        border_style="purple",
        padding=(1, 2),
    )
)
ORCHESTRATOR_SYSTEM_PROMPT = """
Your job is to fullfill user's request step-by-step

# Rule of using different storage

1. the /final_report.md /plan.md /todo-list.md /file-summary.md /plan-hierarchical.md etc. all the files as a general perspective understanding or plan must be stored in the / root level
2. all the input files, are stored in the /workspace and therefore user's reference to user-provided files are located here
3. If the plan involves reading many resources, you should use the root-level storage as temporary storage for your analysis results, summaries, notes, etc., and use the /workspace for user-provided files only.
4. If intended to write enormous files, you should complete part-by-part and store them in the / root level storage, and combine them all together at the end to form the final output file. For example, if you were to write a report including 8 chapters, you should write chapter 1 to chapter 8 separately as /chapter_1.md, /chapter_2.md, ..., /chapter_8.md, and then combine them all together to form the final /final_report.md at the end, to ensure each part's comprehensiveness.

# IMPORTANT NOTE

1. Only WRITE and EXECUTE scripts in the /workspace directory .
2. Do not delegate any file operations to sub-agents, all file operations must be handled by you directly using the provided tools. You can use script for batch operation if needed.
"""
DELEGATION_INSTRUCTIONS = """
ONLY delegate when ABSOLUTELY NECESSARY
1. AFTER doing ASR on audio file, delegate the processing of transcribed results to transcription-processing-agent for further processing, DON'T delegate the processing of audios to subagent!

"""

TRANSCRIPT_POSTPROCESSOR_INSTRUCTIONS = """You are a language post-processing assistant specialized in correcting noisy audio transcriptions and generating structured meeting minutes. 
For context, today's date is {date}.

<Task>
You are given a raw audio transcription that may contain:
- ASR errors (wrong words, homophones, omissions)
- Disfluencies (repetition, fillers, interruptions)
- Incomplete or broken sentences
- Topic jumps and speaker ambiguity

Your task has TWO mandatory outputs:
1. A refined, readable transcript
2. A structured meeting minutes document
</Task>

<Core Principles>
1. **Faithfulness**: Do NOT invent facts, decisions, or participants.
2. **Reconstruction, not summarization** (for transcript refinement).
3. **Abstraction is allowed only in the meeting minutes section.**
4. If content is unclear or ambiguous, explicitly mark it as such.
</Core Principles>

<Processing Steps>
You must internally follow these stages:

Stage 1 — Noise Reduction
- Remove fillers, stutters, repeated fragments
- Correct obvious ASR word errors using context
- Preserve original meaning and intent

Stage 2 — Semantic Reconstruction
- Rebuild broken sentences
- Merge fragmented utterances
- Infer logical sentence boundaries

Stage 3 — Topic Structuring
- Identify major discussion themes
- Detect implicit constraints (budget, manpower, timeline, tools)
- Detect preferences vs. decisions

Stage 4 — Minutes Generation
- Produce a professional, structured meeting summary

Use think_tool between stages if needed to reason about ambiguity or structure.
</Processing Steps>

<Output Format — STRICT>
You must return the following sections IN ORDER:

====================
【一】校订后的完整会议实录（Refined Transcript）
====================
- Continuous prose
- No bullet points
- No commentary

====================
【二】会议纪要（Meeting Minutes）
====================

### 1. 会议背景与目标
### 2. 核心讨论要点
### 3. 明确结论与共识
### 4. 约束条件与风险
### 5. 后续行动项（如有）
### 6. 尚未明确的问题（如有）

If a section has no content, explicitly write “暂无明确内容”。

<Hard Rules>
- Do NOT add new content
- Do NOT change the meaning
- Do NOT use markdown beyond the specified structure
- Do NOT role-play speakers unless clearly identifiable
</Hard Rules>
"""
from datetime import datetime
current_date = datetime.now().strftime("%Y-%m-%d")

# Create meeting transcription  sub-agent
transcription_processing_agent = {
    "name": "transcription-processing-agent",
    "description": "Refine noisy audio transcription and generate structured meeting minutes. No external research.",
    "system_prompt": TRANSCRIPT_POSTPROCESSOR_INSTRUCTIONS.format(date=current_date),
    "tools": [think_tool],
}


_DEFAULT_RENDERER.show_prompt(ORCHESTRATOR_SYSTEM_PROMPT, title="Orchestrator Prompt")
# %%
_DEFAULT_RENDERER.console.print(
    Panel(
        Text("Initialize Deep Agent."),
        title="STEP 5: Initialize Deep Agent",
        border_style="purple",
        padding=(1, 2),
    )
)
from deepagents import create_deep_agent
from langchain.chat_models import init_chat_model
from langchain_openai import ChatOpenAI

# Default model; replace with any LangChain-compatible chat model
# model = init_chat_model(model="anthropic:claude-sonnet-4-5-20250929", temperature=0.0)


def build_model():
    """Pick the default LLM based on environment-driven provider selection."""
    provider = (os.getenv("DEEP_SCHOLAR_LLM_PROVIDER") or "iflow").lower()

    if provider == "deepseek":
        return ChatOpenAI(
            base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1"),
            api_key=os.getenv("DEEPSEEK_API_KEY"),
            model=os.getenv("DEEPSEEK_MODEL", "deepseek-chat"),
            temperature=float(os.getenv("DEEPSEEK_TEMPERATURE", "0.2")),
        )

    if provider == "llama":
        return ChatOpenAI(
            base_url=os.getenv("LLAMA_BASE_URL", "http://localhost:11434/v1"),
            api_key=os.getenv("LLAMA_API_KEY"),
            model=os.getenv("LLAMA_MODEL", "llama3"),
            temperature=float(os.getenv("LLAMA_TEMPERATURE", "0.2")),
        )

    # Default: IFlow (qwen3-max)
    return ChatOpenAI(
        base_url=os.getenv("IFLOW_BASE_URL", "https://apis.iflow.cn/v1"),
        api_key=os.getenv("IFLOW_API_KEY"),
        model=os.getenv("IFLOW_MODEL", "qwen3-max"),
        temperature=float(os.getenv("IFLOW_TEMPERATURE", "0.2")),
    )


model = build_model()

from deepagents import create_deep_agent
from deepagents.backends import CompositeBackend, StateBackend, StoreBackend
from deepagents.backends import FilesystemBackend
from deepagents.backends.protocol import ExecuteResponse, SandboxBackendProtocol
from langgraph.store.memory import InMemoryStore
import subprocess


class LocalSandboxBackend(FilesystemBackend, SandboxBackendProtocol):
    def __init__(
        self,
        *,
        root_dir: str | Path | None = None,
        virtual_mode: bool = True,
        timeout: float = 120.0,
        max_output_bytes: int = 200_000,
        env: dict[str, str] | None = None,
        path_aliases: dict[str, str] | None = None,
    ) -> None:
        super().__init__(root_dir=root_dir, virtual_mode=virtual_mode)
        self._timeout = timeout
        self._max_output_bytes = max_output_bytes
        self._env = env if env is not None else os.environ.copy()
        self._path_aliases = path_aliases or {}

    @property
    def id(self) -> str:
        return f"local:{self.cwd}"

    def _apply_path_aliases(self, command: str) -> str:
        if not self._path_aliases:
            return command
        updated = command
        for virtual_path, real_path in self._path_aliases.items():
            virtual_root = virtual_path.rstrip("/")
            real_root = str(Path(real_path).resolve()).rstrip("/")
            updated = updated.replace(f"{virtual_root}/", f"{real_root}/")
            updated = updated.replace(virtual_root, real_root)
        return updated

    def execute(self, command: str) -> ExecuteResponse:
        if not isinstance(command, str) or not command.strip():
            return ExecuteResponse(
                output="Error: execute expects a non-empty command string.",
                exit_code=1,
                truncated=False,
            )

        command = self._apply_path_aliases(command)
        try:
            result = subprocess.run(
                command,
                shell=True,
                cwd=str(self.cwd),
                capture_output=True,
                text=True,
                timeout=self._timeout,
                env=self._env,
            )
        except subprocess.TimeoutExpired:
            return ExecuteResponse(
                output=f"Error: Command timed out after {self._timeout:.1f} seconds.",
                exit_code=124,
                truncated=False,
            )

        output_parts: list[str] = []
        if result.stdout:
            output_parts.append(result.stdout)
        if result.stderr:
            output_parts.append(result.stderr)
        output = "\n".join(output_parts) if output_parts else "<no output>"

        truncated = False
        if len(output) > self._max_output_bytes:
            output = output[: self._max_output_bytes]
            truncated = True

        return ExecuteResponse(output=output, exit_code=result.returncode, truncated=truncated)


workspace_root = str(Path("./workspace").resolve())
composite_backend = lambda rt: CompositeBackend(
    default=LocalSandboxBackend(
        root_dir=".",
        virtual_mode=True,
        path_aliases={
            "/workspace": workspace_root,
        },
    ),
    routes={
        "/workspace/": LocalSandboxBackend(root_dir="./workspace", virtual_mode=True),
    },
)

agent = create_deep_agent(
    model=model,
    tools=all_tools,
    system_prompt=ORCHESTRATOR_SYSTEM_PROMPT + DELEGATION_INSTRUCTIONS,
    subagents=[transcription_processing_agent],
    backend=composite_backend,
)

# %%

request_dict = {
    "report generation": "Write me a /final_report.md based on the files from the zip file inside the /workspace, write the summary report in pure Chinese, make it extremly long and detailed, use as many as references from Chinese Commnunist Party history or Communism Theory as possible, make it official and academic style, targeting as a report for the central standing committee of the Communist Party of China.",

    "excel_analysis1": "我希望了解给出的excel的整体情况，其中我想知道来自山西省的有哪些人？硕士学历以及更高学历的有哪些人？",
    "meeting_minutes": "我有一个会议的录音文件，我希望生成一份完整详细正规的会议纪要。",
    "try-execute": "Write a script to calculate the fibonacci series and give me the results."
}
request_message = {
    "messages": [
        {
            "role": "user",
            "content": request_dict["try-execute"],
        }
    ],
}

# %%
if __name__ == "__main__":
    for event in agent.stream(request_message):
        _DEFAULT_RENDERER.render_stream_event(event)
# %%
if __name__ == "__main__":
    example_result = agent.invoke(request_message)
# %%
if __name__ == "__main__":
    _DEFAULT_RENDERER.render_final_output(example_result)
    print(example_result['files'])
# %%
