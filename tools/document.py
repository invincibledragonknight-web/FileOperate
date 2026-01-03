from __future__ import annotations

from typing import Any, Dict, List, Mapping, Optional

from langchain.tools import tool

from workspace import resolve_workspace_path

@tool(parse_docstring=True)
def pdf_reader(virtual_pdf_path: str, num_pages: int = 5) -> dict:
    """
    Read the first N pages of a PDF file located inside the workspace.

    This preprocessing function extracts text from the first N pages of a PDF
    file referenced by a virtual workspace path.

    The function operates strictly within the workspace sandbox:
    - Input paths are virtual (e.g. `/workspace/document.pdf`)
    - Execution is performed on resolved real paths
    - Output paths are returned in virtual-path form for agent consumption

    Typical agent usage:
    - Call this when a PDF file is detected during workspace inspection
    - Follow with further operations based on the extracted text

    Args:
        virtual_pdf_path: Virtual path to a PDF file inside the workspace
            (for example `/workspace/document.pdf`). PATH MUST START WITH `/workspace`.
        num_pages: Number of pages to read from the start of the PDF.
    
    Returns:
        A dictionary containing:
        - status: Execution status string
        - pdf: The input virtual PDF path
        - content: Extracted text from the first N pages of the PDF
    
    Raises:
        FileNotFoundError: If the PDF file does not exist.
        ValueError: If the provided path does not point to a PDF file.
    """
    from pypdf import PdfReader

    pdf_path = resolve_workspace_path(virtual_pdf_path)

    if not pdf_path.exists():
        return {
            "status": "error",
            "content": f"File not found: {virtual_pdf_path}"
        }

    if pdf_path.suffix.lower() != ".pdf":
        raise ValueError("Provided file is not a PDF file")

    reader = PdfReader(str(pdf_path))
    content = []

    for i, page in enumerate(reader.pages):
        if i >= num_pages:
            break
        content.append(page.extract_text())

    return {
        "status": "ok",
        "pdf": virtual_pdf_path,
        "content": "\n".join(content),
    }

@tool(parse_docstring=True)
def word_reader(virtual_docx_path: str, max_blocks: int = 200) -> dict:
    """
    Read a Microsoft Word document and convert it to formatted Markdown.

    This preprocessing function extracts structured content from a Word document
    referenced by a virtual workspace path. It returns a Markdown rendering plus
    formatting metadata for paragraphs, runs, and tables.

    The function operates strictly within the workspace sandbox:
    - Input paths are virtual (e.g. `/workspace/document.docx`)
    - Execution is performed on resolved real paths
    - No file content is modified

    Typical agent usage:
    - Call this when a Word document is detected during workspace inspection
    - Use the Markdown output for downstream summarization or extraction
    - Use the formatting metadata when layout fidelity matters

    Args:
        virtual_docx_path: Virtual path to a Word document inside the workspace
            (for example `/workspace/docs/report.docx`).
            PATH MUST START WITH `/workspace`.
        max_blocks: Maximum number of top-level blocks (paragraphs or tables)
            to return. Set to 0 to disable the limit.

    Returns:
        A dictionary containing:
        - status: Execution status string
        - word: The input virtual Word path
        - metadata: Core document metadata (title, author, timestamps)
        - blocks: Structured blocks with formatting info
        - markdown: Markdown rendering of the document
        - blocks_returned: Number of blocks returned
        - truncated: Whether the output was truncated by max_blocks

    Raises:
        ValueError: If the provided path is not a supported Word document.
    """
    try:
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except Exception as exc:
        return {
            "status": "error",
            "word": virtual_docx_path,
            "error": f"python-docx is required to read Word files: {exc}",
        }

    import re

    docx_path = resolve_workspace_path(virtual_docx_path)

    if not docx_path.exists():
        return {
            "status": "error",
            "word": virtual_docx_path,
            "error": f"File not found: {virtual_docx_path}",
        }

    if docx_path.suffix.lower() not in {".docx", ".docm", ".dotx", ".dotm"}:
        raise ValueError("Provided file is not a supported Word document")

    document = Document(str(docx_path))

    def iter_block_items(doc: Document) -> Iterable[Union[Paragraph, Table]]:
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    def escape_markdown_chars(text: str) -> str:
        text = text.replace("\\", "\\\\")
        for ch in ("`", "*", "_", "{", "}", "[", "]", "(", ")", "#", "+", "!", "|", ">"):
            text = text.replace(ch, f"\\{ch}")
        return text

    def escape_html(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def normalize_text(text: str, in_table: bool = False) -> str:
        if not text:
            return ""
        text = text.replace("\r", "")
        text = text.replace("\t", "    ")
        token = "__DOCX_BR__"
        text = text.replace("\n", token)
        text = escape_html(text)
        text = escape_markdown_chars(text)
        if in_table:
            return text.replace(token, "<br>")
        return text.replace(token, "  \n")

    def resolve_emphasis(run) -> Tuple[bool, bool]:
        bold = run.bold
        italic = run.italic
        style_name = (run.style.name if run.style else "").lower()
        if bold is None and ("strong" in style_name or "bold" in style_name):
            bold = True
        if italic is None and ("emphasis" in style_name or "italic" in style_name):
            italic = True
        return bool(bold), bool(italic)

    def run_to_markdown(run, in_table: bool = False) -> str:
        raw_text = run.text
        if raw_text is None or raw_text == "":
            return ""
        text = normalize_text(raw_text, in_table=in_table)
        bold, italic = resolve_emphasis(run)
        strike = True if run.font.strike else False
        if bold and italic:
            text = f"***{text}***"
        elif bold:
            text = f"**{text}**"
        elif italic:
            text = f"*{text}*"
        if strike:
            text = f"~~{text}~~"
        return text

    def paragraph_to_markdown(paragraph: Paragraph, in_table: bool = False) -> str:
        parts: List[str] = []
        for run in paragraph.runs:
            part = run_to_markdown(run, in_table=in_table)
            if part:
                parts.append(part)
        if not parts:
            return normalize_text(paragraph.text or "", in_table=in_table)
        return "".join(parts)

    def detect_heading_level(style_id: Optional[str], style_name: Optional[str]) -> Optional[int]:
        for source in (style_id, style_name):
            if not source:
                continue
            lower = source.lower()
            if lower.startswith("heading"):
                match = re.search(r"(\d+)", source)
                if match:
                    level = int(match.group(1))
                    return max(1, min(level, 6))
        if style_id and style_id.lower() == "title":
            return 1
        if style_name and style_name.lower() == "title":
            return 1
        if style_id and style_id.lower() == "subtitle":
            return 2
        if style_name and style_name.lower() == "subtitle":
            return 2
        return None

    def detect_list_info(paragraph: Paragraph, style_id: Optional[str], style_name: Optional[str]) -> Tuple[Optional[str], Optional[int]]:
        list_type = None
        list_level = None
        combined = f"{style_id or ''} {style_name or ''}".lower()

        if "bullet" in combined:
            list_type = "bullet"
        elif "number" in combined or "decimal" in combined:
            list_type = "number"
        elif "list" in combined:
            list_type = "number"

        p_pr = paragraph._p.pPr
        if p_pr is not None and p_pr.numPr is not None:
            ilvl = p_pr.numPr.ilvl
            if ilvl is not None and ilvl.val is not None:
                list_level = int(ilvl.val) + 1
            if list_type is None:
                list_type = "number"

        if list_level is None and style_name:
            match = re.search(r"(\d+)$", style_name)
            if match:
                list_level = int(match.group(1))

        if list_type and list_level is None:
            list_level = 1

        return list_type, list_level

    def is_quote_style(style_id: Optional[str], style_name: Optional[str]) -> bool:
        combined = f"{style_id or ''} {style_name or ''}".lower()
        return "quote" in combined

    def extract_run_format(run) -> Dict[str, Any]:
        underline = run.underline
        if underline is not None and underline not in (True, False):
            underline = str(underline)

        color = None
        if run.font.color is not None:
            if run.font.color.rgb is not None:
                color = f"#{run.font.color.rgb}"
            elif run.font.color.theme_color is not None:
                color = str(run.font.color.theme_color)

        highlight = None
        if run.font.highlight_color is not None:
            highlight = str(run.font.highlight_color)

        size_pt = run.font.size.pt if run.font.size else None

        return {
            "bold": run.bold,
            "italic": run.italic,
            "underline": underline,
            "strike": run.font.strike,
            "superscript": run.font.superscript,
            "subscript": run.font.subscript,
            "font_name": run.font.name,
            "font_size_pt": size_pt,
            "color": color,
            "highlight": highlight,
            "style": run.style.name if run.style else None,
        }

    def extract_paragraph_info(paragraph: Paragraph) -> Dict[str, Any]:
        style = paragraph.style
        style_name = style.name if style else None
        style_id = style.style_id if style else None
        heading_level = detect_heading_level(style_id, style_name)
        list_type, list_level = detect_list_info(paragraph, style_id, style_name)
        quote = is_quote_style(style_id, style_name)

        block_type = "paragraph"
        if heading_level:
            block_type = "heading"
        elif list_type:
            block_type = "list_item"
        elif quote:
            block_type = "quote"

        runs: List[Dict[str, Any]] = []
        for run in paragraph.runs:
            if run.text is None or run.text == "":
                continue
            runs.append({"text": run.text, "formatting": extract_run_format(run)})

        alignment = None
        if paragraph.alignment is not None:
            alignment = str(paragraph.alignment)

        return {
            "type": block_type,
            "text": paragraph.text or "",
            "style": style_name,
            "style_id": style_id,
            "alignment": alignment,
            "list_type": list_type,
            "list_level": list_level,
            "heading_level": heading_level,
            "runs": runs,
        }

    def table_to_markdown(table: Table) -> Tuple[str, List[List[str]]]:
        rows: List[List[str]] = []
        for row in table.rows:
            row_cells: List[str] = []
            for cell in row.cells:
                paragraphs = []
                for paragraph in cell.paragraphs:
                    text = paragraph_to_markdown(paragraph, in_table=True).strip()
                    if text:
                        paragraphs.append(text)
                cell_text = "<br>".join(paragraphs) if paragraphs else ""
                row_cells.append(cell_text)
            rows.append(row_cells)

        if not rows:
            return "", []

        col_count = max(len(row) for row in rows)
        for row in rows:
            if len(row) < col_count:
                row.extend([""] * (col_count - len(row)))

        header = rows[0]
        separator = ["---"] * col_count
        body_rows = rows[1:]

        def format_row(values: List[str]) -> str:
            return "| " + " | ".join(values) + " |"

        markdown_lines = [
            format_row(header),
            format_row(separator),
        ]
        for row in body_rows:
            markdown_lines.append(format_row(row))
        return "\n".join(markdown_lines), rows

    def extract_table_info(table: Table) -> Dict[str, Any]:
        rows_info: List[List[Dict[str, Any]]] = []
        for row in table.rows:
            row_info: List[Dict[str, Any]] = []
            for cell in row.cells:
                cell_paragraphs = [extract_paragraph_info(p) for p in cell.paragraphs]
                cell_text = "\n".join(p["text"] for p in cell_paragraphs if p["text"])
                row_info.append(
                    {
                        "text": cell_text,
                        "paragraphs": cell_paragraphs,
                    }
                )
            rows_info.append(row_info)

        column_count = 0
        if table.rows:
            column_count = max(len(row.cells) for row in table.rows)

        return {
            "type": "table",
            "row_count": len(table.rows),
            "column_count": column_count,
            "rows": rows_info,
        }

    metadata = {
        "title": document.core_properties.title,
        "subject": document.core_properties.subject,
        "author": document.core_properties.author,
        "created": document.core_properties.created.isoformat() if document.core_properties.created else None,
        "modified": document.core_properties.modified.isoformat() if document.core_properties.modified else None,
        "last_modified_by": document.core_properties.last_modified_by,
        "revision": document.core_properties.revision,
    }

    blocks: List[Dict[str, Any]] = []
    markdown_blocks: List[str] = []
    list_buffer: List[str] = []
    current_list_type: Optional[str] = None
    truncated = False

    def flush_list_buffer() -> None:
        nonlocal current_list_type
        if list_buffer:
            markdown_blocks.append("\n".join(list_buffer))
            list_buffer.clear()
        current_list_type = None

    block_limit = max_blocks if max_blocks and max_blocks > 0 else None

    for block in iter_block_items(document):
        if block_limit is not None and len(blocks) >= block_limit:
            truncated = True
            break

        if isinstance(block, Paragraph):
            info = extract_paragraph_info(block)
            blocks.append(info)

            style = block.style
            style_name = style.name if style else None
            style_id = style.style_id if style else None
            heading_level = detect_heading_level(style_id, style_name)
            list_type, list_level = detect_list_info(block, style_id, style_name)
            is_quote = is_quote_style(style_id, style_name)
            text_md = paragraph_to_markdown(block).strip()

            if list_type:
                if current_list_type and list_type != current_list_type:
                    flush_list_buffer()
                current_list_type = list_type
                indent = "  " * ((list_level or 1) - 1)
                prefix = "- " if list_type == "bullet" else "1. "
                list_buffer.append(f"{indent}{prefix}{text_md}")
                continue

            flush_list_buffer()

            if heading_level:
                markdown_blocks.append(f"{'#' * heading_level} {text_md}".rstrip())
            elif is_quote:
                quote_lines = text_md.splitlines() or [""]
                markdown_blocks.append("\n".join(f"> {line}".rstrip() for line in quote_lines))
            else:
                markdown_blocks.append(text_md)

        elif isinstance(block, Table):
            flush_list_buffer()
            table_markdown, _ = table_to_markdown(block)
            if table_markdown:
                markdown_blocks.append(table_markdown)
            blocks.append(extract_table_info(block))

    flush_list_buffer()

    markdown_output = "\n\n".join(markdown_blocks).strip()

    return {
        "status": "ok",
        "word": virtual_docx_path,
        "metadata": metadata,
        "blocks": blocks,
        "markdown": markdown_output,
        "blocks_returned": len(blocks),
        "truncated": truncated,
    }
